import io
import re
import shutil
from typing import Optional, Tuple

import pdfplumber
import pytesseract
import docx
from app.config import settings
from app.schemas.resume import ResumeError, ResumeUploadItem
from app.services.llm_service import (
    LLMExtractionError,
    _fallback_deterministic_resume,
    extract_structured_resume_with_llm,
)

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".doc", ".docx"}


class ResumeParsingError(Exception):
    def __init__(self, code: str, message: str):
        self.code = code
        self.message = message
        super().__init__(message)


def is_tesseract_available() -> bool:
    if shutil.which("tesseract"):
        return True
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def validate_file(filename: str, content: bytes) -> str:
    ext = "." + filename.split(".")[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeParsingError(
            code="unsupported_file_type",
            message=f"Unsupported file type '{ext}'. Only PDF, TXT, DOC, and DOCX files are accepted.",
        )

    max_bytes = settings.max_file_size_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise ResumeParsingError(
            code="oversized_file",
            message=f"File exceeds maximum allowed size of {settings.max_file_size_mb} MB.",
        )

    if ext == ".pdf" and not content.startswith(b"%PDF"):
        raise ResumeParsingError(
            code="unreadable_pdf",
            message="Invalid or corrupted PDF file header.",
        )

    return ext


def extract_text_from_ocr_page(page) -> str:
    if not is_tesseract_available():
        raise ResumeParsingError(
            code="ocr_unavailable",
            message="Scanned PDF detected, but OCR engine (Tesseract) is not installed or available on this server.",
        )

    try:
        pil_image = page.to_image(resolution=200).original
        text = pytesseract.image_to_string(pil_image) or ""
        return text.strip()
    except ResumeParsingError:
        raise
    except Exception:
        raise ResumeParsingError(
            code="ocr_failed",
            message="Failed to perform OCR on scanned PDF page.",
        )


def extract_pdf_pages_with_fallback(content: bytes) -> Tuple[str, str]:
    extracted_pages = []
    page_methods = []

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            if not pdf.pages:
                raise ResumeParsingError(
                    code="empty_resume",
                    message="PDF file contains no pages.",
                )

            for page in pdf.pages:
                native_text = (page.extract_text() or "").strip()
                if len(native_text) >= 20:
                    extracted_pages.append(native_text)
                    page_methods.append("native")
                else:
                    ocr_text = extract_text_from_ocr_page(page)
                    if len(ocr_text) >= 15:
                        extracted_pages.append(ocr_text)
                        page_methods.append("ocr")
                    else:
                        page_methods.append("failed")

    except ResumeParsingError:
        raise
    except Exception:
        raise ResumeParsingError(
            code="unreadable_pdf",
            message="Could not extract readable text from this PDF.",
        )

    full_text = "\n\n".join(extracted_pages)

    if not page_methods or all(m == "failed" for m in page_methods):
        if "ocr" in page_methods or any(m == "failed" for m in page_methods):
            raise ResumeParsingError(
                code="ocr_failed",
                message="Could not extract readable text from this scanned resume.",
            )
        raise ResumeParsingError(
            code="empty_resume",
            message="No readable text was found in this resume. The file may be scanned or image-based.",
        )

    if all(m == "native" for m in page_methods):
        method = "native"
    elif all(m == "ocr" for m in page_methods):
        method = "ocr"
    else:
        method = "hybrid"

    return full_text, method


def extract_txt_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("latin-1")
        except Exception:
            raise ResumeParsingError(
                code="invalid_text_file",
                message="Could not decode text file content.",
            )


def extract_docx_text(content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        extracted = "\n".join(full_text)
        if extracted.strip():
            return extracted
        return extract_txt_text(content)
    except Exception:
        return extract_txt_text(content)



def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)
    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_email(text: str) -> Optional[str]:
    pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    pattern = r"(?:\+?\d{1,4}[-.\s]?)?(?:\(?\d{2,5}\)?[-.\s]?)?\d{3,5}[-.\s]?\d{3,5}"
    matches = re.finditer(pattern, text)
    for m in matches:
        matched_str = m.group(0).strip()
        digits = re.sub(r"\D", "", matched_str)
        if 10 <= len(digits) <= 15:
            return matched_str
    return None


def parse_resume_file(filename: str, content: bytes) -> ResumeUploadItem:
    try:
        ext = validate_file(filename, content)

        if ext == ".pdf":
            raw_text, method = extract_pdf_pages_with_fallback(content)
        elif ext in {".doc", ".docx"}:
            raw_text = extract_docx_text(content)
            method = "native"
        else:
            raw_text = extract_txt_text(content)
            method = "native"

        normalized = normalize_text(raw_text)

        if len(normalized) < 20:
            if method in {"ocr", "hybrid"}:
                raise ResumeParsingError(
                    code="ocr_failed",
                    message="Could not extract readable text from this scanned resume.",
                )
            raise ResumeParsingError(
                code="empty_resume",
                message="No readable text was found in this resume. The file may be scanned or image-based.",
            )

        det_email = extract_email(normalized)
        det_phone = extract_phone(normalized)

        try:
            structured = extract_structured_resume_with_llm(
                normalized_text=normalized,
                extraction_method=method,
                det_email=det_email,
                det_phone=det_phone,
            )
        except LLMExtractionError:
            structured = _fallback_deterministic_resume(
                normalized_text=normalized,
                extraction_method=method,
                det_email=det_email,
                det_phone=det_phone,
            )

        return ResumeUploadItem(
            filename=filename,
            status="processed",
            extraction_method=method,
            candidate=structured.candidate,
            resume=structured,
            raw_text=raw_text,
            normalized_text=normalized,
        )

    except ResumeParsingError as e:
        return ResumeUploadItem(
            filename=filename,
            status="failed",
            extraction_method=None,
            error=ResumeError(code=e.code, message=e.message),
        )
    except Exception:
        return ResumeUploadItem(
            filename=filename,
            status="failed",
            extraction_method=None,
            error=ResumeError(
                code="extraction_failed",
                message="An unexpected error occurred while processing the resume.",
            ),
        )
